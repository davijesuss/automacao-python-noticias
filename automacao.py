import requests
from bs4 import BeautifulSoup
from docx import Document 

link = "https://g1.globo.com/"

requisicao = requests.get(link)
site = BeautifulSoup(requisicao.text, "html.parser")

titulos_noticias = site.find_all(class_="feed-post-body-title gui-color-primary gui-color-hover")

doc = Document()

for titulo_noticias in titulos_noticias:
    titulo_texto = titulo_noticias.find("p", attrs={"elementtiming": "text-ssr"})
    link_noticia = titulo_noticias.find("a")["href"]
    
    if titulo_texto and link_noticia:

        doc.add_heading(titulo_texto.text, level=1)
        doc.add_paragraph("Link da notícia: " + link_noticia)
        
        requisicao_noticia = requests.get(link_noticia)
        site_noticia = BeautifulSoup(requisicao_noticia.text, "html.parser")
        conteudo_noticia = site_noticia.find_all(class_="content-text__container")
        
        for conteudo in conteudo_noticia:
            doc.add_paragraph(conteudo.text.strip())

        doc.add_paragraph('_______________________________________________________________________________________')
        doc.add_paragraph() 
       
doc.save("noticias_g1.docx")
print("Documento do Word criado com sucesso.")
